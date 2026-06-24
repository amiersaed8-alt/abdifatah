import streamlit as st, pandas as pd, plotly.express as px, io
from docx import Document
from datetime import datetime

st.set_page_config(page_title="Borehole Management", layout="wide", page_icon="🚰")

# --- KORDHINTA XOGTA INTERNETKA (MOCK DATA) ---
if "daily_data" not in st.session_state:
    st.session_state.daily_data = pd.DataFrame([
        {"Date": "2026-06-24", "Station": "Dhamuug", "Borehole": "Dhamuug-01", "Engine_Hours": 8.0, "Solar_Hours": 4.0, "Production_m3": 120.0, "Diesel_Liters": 15.0, "E_Maint": 0.0, "B_Maint": 0.0, "S_Maint": 0.0},
        {"Date": "2026-06-25", "Station": "Afraag", "Borehole": "Afraag-01", "Engine_Hours": 6.0, "Solar_Hours": 6.0, "Production_m3": 95.0, "Diesel_Liters": 12.0, "E_Maint": 50.0, "B_Maint": 0.0, "S_Maint": 20.0}
    ])

df = pd.DataFrame(st.session_state.daily_data)
df['Date'] = pd.to_datetime(df['Date'])
df['Month'] = df['Date'].dt.to_period('M').astype(str)

# Maanta iyo bishan hadda jirta
today_str = datetime.now().strftime("%Y-%m-%d")
this_month_str = datetime.now().strftime("%Y-%m")

df_today = df[df['Date'].dt.strftime('%Y-%m-%d') == today_str]
df_month = df[df['Month'] == this_month_str]

# --- LOGO IYO CINWAANKA ---
col_l, col_t = st.columns([1, 4])
with col_l:
    logo = st.file_uploader("Logo", type=["png", "jpg"], label_visibility="collapsed")
    if logo: st.image(logo, width=100)
with col_t:
    st.title("🚰 Water Supply & Borehole Management Dashboard")

st.markdown("---")

t1, t2, t3 = st.tabs(["📊 Dashboard-ka Maamulka", "📝 Entry Form", "📄 Word Reports"])

# ==================== TAB 1: MAAMULKA ====================
with t1:
    # --- 1. MAAMULKA DEGDEGGE AH (QUICK INSIGHTS) ---
    st.subheader("🎯 Maamulka Degdegga ah (Quick Insights)")
    i1, i2, i3, i4, i5 = st.columns(5)
    
    if not df.empty:
        # Ceelka ugu wax-soo-saarka badan
        top_bh = df.groupby('Borehole')['Production_m3'].sum().idxmax()
        i1.metric("Ceelka ugu Product badan", top_bh)
        
        # Engine-ka ugu saacadaha badan
        top_eng = df.groupby('Borehole')['Engine_Hours'].sum().idxmax()
        i2.metric("Engine-ka ugu Saacadaha badan", top_eng)
        
        # Solar-ka ugu shaqada badan
        top_sol = df.groupby('Borehole')['Solar_Hours'].sum().idxmax()
        i3.metric("Solar-ka ugu Shaqada badan", top_sol)
        
        # Ceelka ugu shidaalka badan
        top_dsl = df.groupby('Borehole')['Diesel_Liters'].sum().idxmax()
        i4.metric("Ceelka ugu Shidaalka badan", top_dsl)
        
        # Ceelka ugu maintenance-ka badan bishii
        df['Total_Maint'] = df['E_Maint'] + df['B_Maint'] + df['S_Maint']
        top_maint = df.groupby('Borehole')['Total_Maint'].sum().idxmax()
        i5.metric("Ceelka ugu Maintenance badan", top_maint)
    else:
        for i in [i1, i2, i3, i4, i5]: i.metric("-", "Xog la'aan")

    st.markdown("---")
    
    # --- 2. OPERATIONS ---
    st.subheader("⚙️ Operations (Wax-soo-saarka & Isticmaalka)")
    o1, o2, o3, o4 = st.columns(4)
    o1.metric("Daily Production (m³)", f"{df_today['Production_m3'].sum():,}")
    o1.metric("Monthly Production (m³)", f"{df_month['Production_m3'].sum():,}")
    
    o2.metric("Daily Diesel Total (L)", f"{df_today['Diesel_Liters'].sum():,}")
    o2.metric("Monthly Diesel Total (L)", f"{df_month['Diesel_Liters'].sum():,}")
    
    o3.metric("Daily Engine Hours", f"{df_today['Engine_Hours'].sum():,}")
    o3.metric("Monthly Engine Hours", f"{df_month['Engine_Hours'].sum():,}")
    
    o4.metric("Daily Solar Hours", f"{df_today['Solar_Hours'].sum():,}")
    o4.metric("Monthly Solar Hours", f"{df_month['Solar_Hours'].sum():,}")

    st.markdown("---")

    # --- 3. COMPARISONS ---
    st.subheader("⚖️ Comparisons (Isbarbardhigga)")
    c1, c2 = st.columns(2)
    with c1:fig_bh = px.bar(df, x="Borehole", y="Production_m3", color="Station", title="Production by Borehole & Station")
        st.plotly_chart(fig_bh, use_container_width=True)
    with c2:
        st.write("Engine vs Solar Hours")
        df_melted = df.melt(id_vars=["Borehole"], value_vars=["Engine_Hours", "Solar_Hours"], var_name="Type", value_name="Hours")
        fig_eng = px.bar(df_melted, x="Borehole", y="Hours", color="Type", barmode="group", title="Engine vs Solar Hours")
        st.plotly_chart(fig_eng, use_container_width=True)

    st.markdown("---")

    # --- 4. MAINTENANCE ---
    st.subheader("🔧 Maintenance (Dayactirka)")
    m1, m2, m3, m4 = st.columns(4)
    m1.metric("Engine Maintenance Cost", f"${df['E_Maint'].sum():,}")
    m2.metric("Borehole/Pump Cost", f"${df['B_Maint'].sum():,}")
    m3.metric("Solar Maintenance Cost", f"${df['S_Maint'].sum():,}")
    
    total_m_cost = df_month['E_Maint'].sum() + df_month['B_Maint'].sum() + df_month['S_Maint'].sum()
    m4.metric("Monthly Maint Cost (Total)", f"${total_m_cost:,}", delta_color="inverse")

# ==================== TAB 2: ENTRY FORM ====================
with t2:
    st.header("📝 Foomka Gelinta Xogta Maalinlahay")
    with st.form("entry_form", clear_on_submit=True):
        col1, col2 = st.columns(2)
        with col1:
            dt = st.date_input("Date", datetime.now())
            stn = st.selectbox("Select Station", ["Dhamuug", "Afraag"])
            bh = st.text_input("Qor Magaca Ceelka (Borehole)", "Dhamuug-01")
            e1 = st.number_input("Engine Hours", min_value=0.0, step=0.5)
            sl = st.number_input("Solar Hours", min_value=0.0, step=0.5)
        with col2:
            pr = st.number_input("Production (m³)", min_value=0.0, step=10.0)
            ds = st.number_input("Diesel Consumed (Liters)", min_value=0.0, step=1.0)
            em = st.number_input("Engine Maint Cost ($)", min_value=0.0, step=5.0)
            bm = st.number_input("Borehole/Pump Maint Cost ($)", min_value=0.0, step=5.0)
            sm = st.number_input("Solar Maint Cost ($)", min_value=0.0, step=5.0)
            
        if st.form_submit_button("Save Entry"):
            new_row = {"Date": pd.to_datetime(dt), "Station": stn, "Borehole": bh, "Engine_Hours": e1, "Solar_Hours": sl, "Production_m3": pr, "Diesel_Liters": ds, "E_Maint": em, "B_Maint": bm, "S_Maint": sm}
            st.session_state.daily_data = pd.concat([st.session_state.daily_data, pd.DataFrame([new_row])], ignore_index=True)
            st.success("Xogtii waa la kaydiyey!"); st.rerun()
            
    st.subheader("Shaxda Xogta Guud")
    st.dataframe(st.session_state.daily_data, use_container_width=True)

# ==================== TAB 3: EXPORT ====================
with t3:
    st.header("📄 Export Word Report")
    if st.button("Generate Word Report"):
        doc = Document()
        doc.add_heading('Water Supply Executive Report', level=1)
        table = doc.add_table(rows=1, cols=6)
        for i, name in enumerate(['Date', 'Borehole', 'Prod (m³)', 'Diesel (L)', 'Eng Hours', 'Maint Cost']):
            table.rows.cells[i].text = name
        for _, r in st.session_state.daily_data.iterrows():
            cells = table.add_row().cells
            cells.text, cells.text, cells.text, cells.text, cells.text, cells.text = str(r['Date'])[:10], str(r['Borehole']), str(r['Production_m3']), str(r['Diesel_Liters']), str(r['Engine_Hours']), f"${r['E_Maint']+r['B_Maint']+r['S_Maint']}"
        bio = io.BytesIO(); doc.save(bio)
        st.download_button(label="📥 Download Word", data=bio.getvalue(), file_name="Executive_Report.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        st.write("Borehole vs Borehole & Station vs Station")
